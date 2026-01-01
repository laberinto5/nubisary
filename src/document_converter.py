"""Document conversion module for PDF, DOC, and DOCX files."""

from typing import Optional
import os
import re

from src.validators import FileValidationError


class DocumentConversionError(Exception):
    """Base exception for document conversion errors."""
    pass


class UnsupportedFormatError(DocumentConversionError):
    """Exception raised when document format is not supported."""
    pass


def is_pdf_file(file_path: str) -> bool:
    """Check if file is a PDF based on extension."""
    return file_path.lower().endswith('.pdf')


def is_doc_file(file_path: str) -> bool:
    """Check if file is a DOC file based on extension."""
    return file_path.lower().endswith('.doc')


def is_docx_file(file_path: str) -> bool:
    """Check if file is a DOCX file based on extension."""
    return file_path.lower().endswith('.docx')


def is_convertible_document(file_path: str) -> bool:
    """Check if file is a convertible document (PDF, DOC, DOCX)."""
    return is_pdf_file(file_path) or is_doc_file(file_path) or is_docx_file(file_path)


def clean_converted_text(text: str) -> str:
    """
    Clean converted text by removing common artifacts from PDF/DOC conversion.
    
    Performs the following cleaning operations:
    1. Removes excessive blank lines (more than one consecutive newline)
    2. Removes lines that appear to be page numbers (only numbers, optionally with dashes/parentheses)
    3. Removes lines that appear to be Roman numerals (I, II, III, IV, etc., optionally with symbols)
    
    Args:
        text: Raw converted text
        
    Returns:
        Cleaned text
    """
    lines = text.split('\n')
    cleaned_lines = []
    prev_line_empty = False
    
    # Pattern to match page numbers: lines with only digits, optionally with dashes, parentheses, brackets, or spaces
    # Examples: "1", "1-", "-1", "(1)", "[1]", "1  ", "  1  ", etc.
    page_number_pattern = re.compile(r'^[\s\-\(\)\[\]]*\d+[\s\-\(\)\[\]]*$')
    
    # Pattern to match Roman numerals: valid Roman numeral characters (I, V, X, L, C, D, M)
    # optionally with dashes, parentheses, brackets, or spaces
    # Examples: "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "[II]", "-III-", "(IV)", etc.
    # Valid Roman numeral pattern: matches sequences of I, V, X, L, C, D, M in valid combinations
    roman_numeral_pattern = re.compile(
        r'^[\s\-\(\)\[\]]*'  # Optional leading symbols/spaces
        r'[IVXLCDM]+'  # One or more valid Roman numeral characters
        r'[\s\-\(\)\[\]]*$',  # Optional trailing symbols/spaces
        re.IGNORECASE
    )
    
    def is_valid_roman_numeral(text: str) -> bool:
        """
        Check if a string is a valid Roman numeral.
        This is a basic check - more sophisticated validation could be added.
        """
        # Remove symbols and spaces
        cleaned = re.sub(r'[\s\-\(\)\[\]]', '', text).upper()
        
        # Check if it contains only valid Roman numeral characters
        if not re.match(r'^[IVXLCDM]+$', cleaned):
            return False
        
        # Basic validation: check for invalid patterns
        # (This is simplified - full Roman numeral validation is more complex)
        # For now, accept any sequence of valid characters
        # Common invalid patterns: IIII, VV, etc. (but we'll be lenient)
        return len(cleaned) > 0
    
    for line in lines:
        # Check if line is empty
        is_empty = not line.strip()
        
        # Skip if it's a page number (only numbers with optional formatting)
        if page_number_pattern.match(line.strip()):
            continue
        
        # Skip if it's a Roman numeral (only Roman numerals with optional formatting)
        stripped_line = line.strip()
        if roman_numeral_pattern.match(stripped_line) and is_valid_roman_numeral(stripped_line):
            continue
        
        # Handle multiple consecutive blank lines - keep only one
        if is_empty:
            if not prev_line_empty:
                cleaned_lines.append('')
            prev_line_empty = True
        else:
            cleaned_lines.append(line)
            prev_line_empty = False
    
    # Remove trailing blank lines
    while cleaned_lines and not cleaned_lines[-1].strip():
        cleaned_lines.pop()
    
    return '\n'.join(cleaned_lines)


def convert_pdf_to_text(pdf_path: str, clean_text: bool = True) -> str:
    """
    Convert PDF file to plain text.
    
    Args:
        pdf_path: Path to PDF file
        clean_text: If True, apply text cleaning (remove page numbers, excessive blank lines)
        
    Returns:
        Extracted text content
        
    Raises:
        FileValidationError: If file doesn't exist or can't be read
        DocumentConversionError: If conversion fails
    """
    try:
        import pypdf
    except ImportError:
        raise DocumentConversionError(
            'pypdf library is required for PDF conversion. '
            'Install it with: pip install pypdf'
        )
    
    if not os.path.isfile(pdf_path):
        raise FileValidationError(f'PDF file does not exist: {pdf_path}')
    
    try:
        text_content = []
        with open(pdf_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                raise DocumentConversionError(
                    f'PDF file is encrypted/password-protected: {pdf_path}. '
                    'Please decrypt the PDF first.'
                )
            
            # Extract text from all pages
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                except Exception as e:
                    raise DocumentConversionError(
                        f'Error extracting text from page {page_num} of {pdf_path}: {e}'
                    )
        
        if not text_content:
            raise DocumentConversionError(
                f'No text could be extracted from PDF: {pdf_path}. '
                'The PDF might contain only images or be corrupted.'
            )
        
        result = '\n\n'.join(text_content)
        
        # Apply text cleaning if requested
        if clean_text:
            result = clean_converted_text(result)
        
        return result
        
    except pypdf.errors.PdfReadError as e:
        raise DocumentConversionError(f'Invalid or corrupted PDF file {pdf_path}: {e}')
    except Exception as e:
        if isinstance(e, DocumentConversionError):
            raise
        raise DocumentConversionError(f'Error converting PDF {pdf_path}: {e}')


def convert_docx_to_text(docx_path: str, clean_text: bool = True) -> str:
    """
    Convert DOCX file to plain text.
    
    Args:
        docx_path: Path to DOCX file
        clean_text: If True, apply text cleaning (remove page numbers, excessive blank lines)
        
    Returns:
        Extracted text content
        
    Raises:
        FileValidationError: If file doesn't exist or can't be read
        DocumentConversionError: If conversion fails
    """
    try:
        from docx import Document
    except ImportError:
        raise DocumentConversionError(
            'python-docx library is required for DOCX conversion. '
            'Install it with: pip install python-docx'
        )
    
    if not os.path.isfile(docx_path):
        raise FileValidationError(f'DOCX file does not exist: {docx_path}')
    
    try:
        doc = Document(docx_path)
        text_content = []
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        if not text_content:
            raise DocumentConversionError(
                f'No text could be extracted from DOCX: {docx_path}. '
                'The document might be empty or contain only images.'
            )
        
        result = '\n\n'.join(text_content)
        
        # Apply text cleaning if requested
        if clean_text:
            result = clean_converted_text(result)
        
        return result
        
    except Exception as e:
        if isinstance(e, DocumentConversionError):
            raise
        raise DocumentConversionError(f'Error converting DOCX {docx_path}: {e}')


def convert_doc_to_text(doc_path: str, clean_text: bool = True) -> str:
    """
    Convert DOC file to plain text.
    
    Note: DOC format (older Word format) is a proprietary binary format that cannot be
    reliably converted using pure Python libraries without external dependencies.
    
    This function raises an error with clear instructions to convert the file first.
    
    Args:
        doc_path: Path to DOC file
        clean_text: If True, apply text cleaning (remove page numbers, excessive blank lines)
        
    Returns:
        Extracted text content (never reached, always raises error)
        
    Raises:
        FileValidationError: If file doesn't exist or can't be read
        DocumentConversionError: Always raised with conversion instructions
    """
    if not os.path.isfile(doc_path):
        raise FileValidationError(f'DOC file does not exist: {doc_path}')
    
    raise DocumentConversionError(
        f'DOC file format is not directly supported.\n\n'
        f'The .doc format (older Microsoft Word format) is a proprietary binary format '
        f'that requires external tools to convert. To avoid adding system dependencies, '
        f'please convert your .doc file to a supported format first.\n\n'
        f'Recommended solutions:\n'
        f'1. Convert DOC to DOCX: Open the file in Microsoft Word or LibreOffice and save as .docx\n'
        f'2. Convert DOC to TXT: Open the file and save as plain text (.txt)\n'
        f'3. Use LibreOffice command line (if available):\n'
        f'   libreoffice --headless --convert-to docx "{doc_path}"\n\n'
        f'Once converted, you can use this tool with the .docx or .txt file.'
    )


def convert_document_to_text(file_path: str, clean_text: bool = True) -> str:
    """
    Convert a document (PDF, DOC, DOCX) to plain text.
    Auto-detects the file type and uses the appropriate converter.
    
    Args:
        file_path: Path to document file
        clean_text: If True, apply text cleaning (remove page numbers, excessive blank lines)
        
    Returns:
        Extracted text content
        
    Raises:
        UnsupportedFormatError: If file format is not supported
        DocumentConversionError: If conversion fails
    """
    if is_pdf_file(file_path):
        return convert_pdf_to_text(file_path, clean_text=clean_text)
    elif is_docx_file(file_path):
        return convert_docx_to_text(file_path, clean_text=clean_text)
    elif is_doc_file(file_path):
        return convert_doc_to_text(file_path, clean_text=clean_text)
    else:
        raise UnsupportedFormatError(
            f'Unsupported document format: {file_path}. '
            'Supported formats: PDF, DOC, DOCX'
        )

